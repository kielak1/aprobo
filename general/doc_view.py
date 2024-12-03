from django.http import HttpResponse
from docx import Document
from general.models import Meeting, Resolution


def protokol_rady_architektury(request, id_rady):
    # Pobierz dane spotkania na podstawie id_rady
    meeting = Meeting.objects.get(id=id_rady)

    # Tworzenie nowego dokumentu Word
    document = Document()

    # Dodaj tytuł dokumentu
    document.add_heading(
        f"Protokół z posiedzenia Rady Architektury w dniu {meeting.meeting_date}", 0
    )

    # Lista zaproszonych
    document.add_heading("Lista zaproszonych:", level=1)
    for invited in meeting.invitations.all():
        document.add_paragraph(
            f"{invited.first_name} {invited.last_name} ({invited.email})"
        )

    # Lista obecnych
    document.add_heading("Lista obecnych:", level=1)
    for attendee in meeting.attendance_list.all():
        document.add_paragraph(
            f"{attendee.first_name} {attendee.last_name} ({attendee.email})"
        )

    # Generalne postanowienia
    document.add_heading("Postanowienia ogólne:", level=1)
    if meeting.general_resolution:
        document.add_paragraph(meeting.general_resolution)
    else:
        document.add_paragraph("Ogólne postanowienia nie zostały podjęte.")

    # Omawiane Pomysły i Potrzeby

    # Pobierz postanowienia z powiązanymi ideami i potrzebami
    resolutions = Resolution.objects.filter(meeting=meeting)

    document.add_heading("Omawiane Pomysły:", level=1)
    for idx, resolution in enumerate(resolutions, start=1):
        if resolution.idea:
            subject = (
                resolution.idea.subject if resolution.idea else resolution.need.subject
            )
            section = (
                resolution.idea.section if resolution.idea else resolution.need.section
            )
            number = resolution.idea.id if resolution.idea.id else resolution.need.id
            data_utworzenia = (
                resolution.idea.data_utworzenia
                if resolution.idea
                else resolution.need.data_utworzenia
            )
            prowadzący = (
                resolution.idea.osoba_prowadzaca
                if resolution.idea
                else resolution.need.osoba_prowadzaca
            )
            orientacyjny_budzet = (
                resolution.idea.orientacynjy_budzet
                if resolution.idea
                else resolution.need.orientacynjy_budzet
            )

            document.add_heading(f"{number}. {subject}", level=2)
            document.add_paragraph(
                f"Dział: {section}, Koordynator IT: {prowadzący.get_full_name()}, data utworzenia: {data_utworzenia}, budżet: {orientacyjny_budzet}"
            )
            document.add_paragraph(f"Podjęte ustalenia: {resolution.resolution_text}")

    document.add_heading("Omawiane Potrzeby:", level=1)
    for idx, resolution in enumerate(resolutions, start=1):
        if resolution.need:
            subject = (
                resolution.idea.subject if resolution.idea else resolution.need.subject
            )
            section = (
                resolution.idea.section if resolution.idea else resolution.need.section
            )
            number = resolution.idea.id if resolution.idea else resolution.need.id
            data_utworzenia = (
                resolution.idea.data_utworzenia
                if resolution.idea
                else resolution.need.data_utworzenia
            )
            prowadzący = (
                resolution.idea.osoba_prowadzaca
                if resolution.idea
                else resolution.need.osoba_prowadzaca
            )
            orientacyjny_budzet = (
                resolution.idea.orientacynjy_budzet
                if resolution.idea
                else resolution.need.orientacynjy_budzet
            )

            document.add_heading(f"{number}. {subject}", level=2)
            document.add_paragraph(
                f"Dział: {section}, Koordynator IT: {prowadzący.get_full_name()}, data utworzenia: {data_utworzenia}, budżet: {orientacyjny_budzet}"
            )
            document.add_paragraph(f"Podjęte ustalenia: {resolution.resolution_text}")

    # Zapisz dokument do obiektu HttpResponse
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        f"attachment; filename=protokol_rady_{id_rady}.docx"
    )

    # Zapisz dokument do strumienia odpowiedzi
    document.save(response)

    return response
